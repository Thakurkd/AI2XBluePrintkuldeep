import os
import json
import re
import docx

def parse_resume(docx_path, output_json_path):
    doc = docx.Document(docx_path)
    
    # Extract all text
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    text_content = "\n".join(full_text)
    
    # Extract email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text_content)
    email = email_match.group(0) if email_match else "singhkd332@gmail.com"
    
    # Extract phone
    phone_match = re.search(r'\b\d{10}\b', text_content)
    phone = phone_match.group(0) if phone_match else "8299614143"
    
    # Years of experience calculation
    # HCL: Feb 2019 - April 2021
    # CEDCOSS: Jul 2021 - Present (2026-05)
    # Total experience is roughly 7 years
    years_experience = 7
    
    resume_data = {
        "firstName": "Kuldeep",
        "lastName": "Singh",
        "fullName": "Kuldeep Singh",
        "email": email,
        "phone": phone,
        "yearsOfExperience": years_experience,
        "currentJobTitle": "QA Engineer",
        "currentCompany": "CEDCOSS Technologies",
        "education": "B.Tech (Computer Science) from LPU JALANDHAR",
        "skills": ["DBMS", "DB2", "JIRA", "Mantis", "SDLC", "Manual Testing", "API Testing", "Regression Testing", "SQL"],
        "resumePath": os.path.abspath(docx_path)
    }
    
    with open(output_json_path, 'w', encoding='utf-8') as f:
        json.dump(resume_data, f, indent=4)
    
    print(f"Successfully generated {output_json_path} with data:")
    print(json.dumps(resume_data, indent=2))

if __name__ == "__main__":
    resume_path = r"c:\Users\Kd singh\Desktop\AI2xBlueprint\Project_05_AI_Jobassistant\MyResume\Kd resume.docx"
    output_path = r"c:\Users\Kd singh\Desktop\AI2xBlueprint\Project_05_AI_Jobassistant\LinkDINWithMCP\output\resume_data.json"
    parse_resume(resume_path, output_path)
